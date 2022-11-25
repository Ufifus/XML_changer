import docx

path_to_testfile = '/home/artem/Downloads/Выписная пневмония 11.docx'

doc = docx.Document(docx=path_to_testfile)

print(len(doc.paragraphs))

full_text = []
for par in doc.paragraphs:
    full_text.append(par.text)

text = '\n'.join(full_text)
print(text)